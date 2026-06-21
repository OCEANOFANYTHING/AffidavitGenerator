import os
from datetime import datetime
import customtkinter as ctk
from docx import Document
from docx.shared import Pt
from tkinter import messagebox
from tkcalendar import DateEntry

# ---------- Core Generator Logic ----------
def replace_text_in_paragraph(paragraph, placeholder, value, font_size=12):
    """Replace placeholder in paragraph while preserving formatting and applying Bookman Old Style"""
    if placeholder in paragraph.text:
        # Store original paragraph formatting
        original_alignment = paragraph.alignment
        original_style = paragraph.style

        # Replace the text
        full_text = paragraph.text
        new_text = full_text.replace(placeholder, value)

        # Clear existing runs
        paragraph.clear()

        # Add new text as a single run with Bookman Old Style font
        run = paragraph.add_run(new_text)
        run.font.name = 'Bookman Old Style'
        run.font.size = Pt(font_size)

        # Restore paragraph formatting
        paragraph.alignment = original_alignment
        paragraph.style = original_style

def generate_affidavits(data):
    templates = ["Shedule-1C.docx", "Self-Declaration.docx", "Introducer-CR.docx"]
    output_dir = os.path.join("output", data["Applicant Name"])
    os.makedirs(output_dir, exist_ok=True)

    # Build addresses from components
    indian_address = ", ".join([
        data.get("IND_Village", ""),
        data.get("IND_Post_Office", ""),
        data.get("IND_Police_Station", ""),
        data.get("IND_District", ""),
        data.get("IND_Pin_Code", "")
    ])

    bangladesh_address = ", ".join([
        data.get("BD_Village", ""),
        data.get("BD_Post_Office", ""),
        data.get("BD_Police_Station", ""),
        data.get("BD_District", "")
    ])

    introducer_address = ", ".join([
        data.get("INT_Village", ""),
        data.get("INT_Post_Office", ""),
        data.get("INT_Police_Station", ""),
        data.get("INT_District", ""),
        data.get("INT_Pin_Code", "")
    ])

    # Map form fields to template placeholders
    placeholder_map = {
        "{{NAME}}": data.get("Applicant Name", ""),
        "{{FATHER_NAME}}": data.get("Applicants Father Name", ""),
        "{{IND_ADDRESS}}": indian_address,
        "{{BD_ADDRESS}}": bangladesh_address,
        "{{DOE}}": data.get("Date of Entry", ""),
        "{{DATE}}": data.get("Current Date (Auto Fetch)", ""),
        "{{CR_PROVIDER_NAME}}": data.get("Introducer Name", ""),
        "{{CR_PROVIDER_AGE}}": data.get("Introducer Age", ""),
        "{{CR_PROVIDER_OCCUPATION}}": data.get("Introducer Occupation", ""),
        "{{CR_PROVIDER_FATHER_NAME}}": data.get("Introducer Father Name", ""),
        "{{CR_PROVIDER_ADDRESS}}": introducer_address,
    }

    for template in templates:
        path = os.path.join("affidavit", template)
        if not os.path.exists(path):
            print(f"Template not found: {template}")
            continue

        # Set font size based on template: Shedule-1C uses 9.5pt, others use 12pt
        font_size = 9.5 if template == "Shedule-1C.docx" else 12

        doc = Document(path)

        # Replace in paragraphs with font formatting
        for para in doc.paragraphs:
            for placeholder, value in placeholder_map.items():
                replace_text_in_paragraph(para, placeholder, value, font_size)

        # Replace in tables with font formatting
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in placeholder_map.items():
                            replace_text_in_paragraph(para, placeholder, value, font_size)

        output_path = os.path.join(output_dir, template)
        doc.save(output_path)

    return True

# ---------- GUI ----------
def submit_form():
    # Validate required fields
    data = {}
    
    # Basic text fields
    basic_fields = ["Applicant Name", "Applicants Father Name", "Introducer Name", 
                    "Introducer Occupation", "Introducer Father Name"]
    
    for field in basic_fields:
        widget = entries[field]
        value = widget.get().strip()
        if not value:
            messagebox.showerror("Error", f"Please fill in: {field}")
            return
        data[field] = value
    
    # Indian Address fields
    indian_address_fields = [
        ("IND_Village", "Indian Address - Village"),
        ("IND_Post_Office", "Indian Address - Post Office"),
        ("IND_Police_Station", "Indian Address - Police Station"),
        ("IND_District", "Indian Address - District"),
        ("IND_Pin_Code", "Indian Address - Pin Code")
    ]
    
    for field_key, field_label in indian_address_fields:
        widget = entries[field_key]
        value = widget.get().strip()
        if not value:
            messagebox.showerror("Error", f"Please fill in: {field_label}")
            return
        data[field_key] = value
    
    # Bangladesh Address fields
    bangladesh_address_fields = [
        ("BD_Village", "Bangladesh Address - Village"),
        ("BD_Post_Office", "Bangladesh Address - Post Office"),
        ("BD_Police_Station", "Bangladesh Address - Police Station"),
        ("BD_District", "Bangladesh Address - District")
    ]
    
    for field_key, field_label in bangladesh_address_fields:
        widget = entries[field_key]
        value = widget.get().strip()
        if not value:
            messagebox.showerror("Error", f"Please fill in: {field_label}")
            return
        data[field_key] = value
    
    # Introducer Address fields
    introducer_address_fields = [
        ("INT_Village", "Introducer Address - Village"),
        ("INT_Post_Office", "Introducer Address - Post Office"),
        ("INT_Police_Station", "Introducer Address - Police Station"),
        ("INT_District", "Introducer Address - District"),
        ("INT_Pin_Code", "Introducer Address - Pin Code")
    ]
    
    for field_key, field_label in introducer_address_fields:
        widget = entries[field_key]
        value = widget.get().strip()
        if not value:
            messagebox.showerror("Error", f"Please fill in: {field_label}")
            return
        data[field_key] = value
    
    # Get date of entry
    date_entry_value = date_entry_widget.get_date()
    data["Date of Entry"] = date_entry_value.strftime("%d-%m-%Y")
    
    # Get introducer age
    age_value = age_entry.get().strip()
    if not age_value:
        messagebox.showerror("Error", "Please fill in: Introducer Age")
        return
    try:
        age_int = int(age_value)
        if age_int < 18 or age_int > 100:
            messagebox.showerror("Error", "Introducer Age must be between 18 and 100")
            return
        data["Introducer Age"] = str(age_int)
    except ValueError:
        messagebox.showerror("Error", "Introducer Age must be a valid number")
        return
    
    # Auto-generate current date
    data["Current Date (Auto Fetch)"] = datetime.now().strftime("%d-%m-%Y")
    
    try:
        success = generate_affidavits(data)
        if success:
            messagebox.showinfo(
                "Success", 
                f"✅ Affidavits generated successfully!\n\nSaved to: output/{data['Applicant Name']}/\n\nFiles:\n• Shedule-1C.docx\n• Self-Declaration.docx\n• Introducer-CR.docx"
            )
            status_label.configure(
                text=f"✅ Generated successfully in 'output/{data['Applicant Name']}'", 
                text_color="green"
            )
    except Exception as e:
        messagebox.showerror("Error", f"Failed to generate affidavits:\n{str(e)}")
        status_label.configure(text="❌ Generation failed", text_color="red")

def clear_form():
    """Clear all form fields"""
    for widget in entries.values():
        if isinstance(widget, ctk.CTkTextbox):
            widget.delete("1.0", "end")
        else:
            widget.delete(0, 'end')
    age_entry.delete(0, 'end')
    date_entry_widget.set_date(datetime.now())
    status_label.configure(text="", text_color="gray")

def focus_next_widget(event):
    """Move focus to next widget when Tab is pressed"""
    next_widget = event.widget.tk_focusNext()
    next_widget.focus()
    scroll_to_widget(next_widget)
    return "break"  # Prevent default Tab behavior

# Store references to introducer address entries for enabling/disabling
introducer_address_entries = []

def toggle_same_address():
    """Toggle introducer address fields based on checkbox state"""
    is_checked = same_address_var.get()

    if is_checked:
        # Copy Indian Address values to Introducer Address fields
        for field_key in ["INT_Village", "INT_Post_Office", "INT_Police_Station", "INT_District", "INT_Pin_Code"]:
            if field_key in entries:
                # Get corresponding Indian address field
                ind_key = field_key.replace("INT_", "IND_")
                ind_value = entries.get(ind_key, None)
                if ind_value:
                    entries[field_key].delete(0, 'end')
                    entries[field_key].insert(0, ind_value.get().strip())

        # Disable introducer address fields
        for entry in introducer_address_entries:
            entry.configure(state="disabled")
    else:
        # Enable introducer address fields
        for entry in introducer_address_entries:
            entry.configure(state="normal")

def scroll_to_widget(widget):
    """Scroll the scrollable frame to make the widget visible"""
    try:
        canvas_see(widget)
    except Exception:
        pass

def canvas_see(widget):
    """Scroll canvas to make widget visible using screen coordinates"""
    try:
        canvas = frame._parent_canvas
        canvas.update_idletasks()

        # Get widget's screen position
        widget.update_idletasks()
        widget_top = widget.winfo_rooty()
        widget_bottom = widget_top + widget.winfo_height()

        # Get canvas's visible screen area
        canvas_top = canvas.winfo_rooty()
        canvas_bottom = canvas_top + canvas.winfo_height()

        # Calculate total scrollable height
        scrollregion = canvas.cget('scrollregion')
        if scrollregion:
            _, _, _, total_height = map(float, scrollregion.split())
        else:
            total_height = canvas.winfo_height()

        # Determine if scrolling is needed
        if widget_top < canvas_top + 60:
            # Widget is above visible area - scroll up
            # Find where the widget is in the scrollable content
            # The widget's y position relative to the interior frame
            interior = None
            for child in canvas.winfo_children():
                if child.winfo_name() and 'scrollbar' not in child.winfo_name().lower():
                    interior = child
                    break

            if interior:
                interior_y = interior.winfo_rooty()
                widget_rel_y = widget_top - interior_y
                scroll_fraction = widget_rel_y / total_height if total_height > 0 else 0
                canvas.yview_moveto(max(0, scroll_fraction - 0.03))
        elif widget_bottom > canvas_bottom - 60:
            # Widget is below visible area - scroll down
            interior = None
            for child in canvas.winfo_children():
                if child.winfo_name() and 'scrollbar' not in child.winfo_name().lower():
                    interior = child
                    break

            if interior:
                interior_y = interior.winfo_rooty()
                widget_rel_y = widget_top - interior_y
                scroll_fraction = widget_rel_y / total_height if total_height > 0 else 0
                # Leave some space at top (15% of visible area)
                visible_fraction = canvas.winfo_height() / total_height if total_height > 0 else 0.5
                canvas.yview_moveto(max(0, scroll_fraction - visible_fraction + 0.1))
    except Exception as e:
        print(f"Scroll error: {e}")

def on_widget_focus(event):
    """Scroll to widget when it receives focus via any method (Tab, click, etc.)"""
    widget = event.widget
    # Only scroll for actual entry widgets, not labels or other widgets
    if isinstance(widget, (ctk.CTkEntry, ctk.CTkTextbox)):
        scroll_to_widget(widget)

def validate_number_input(value):
    """Validate that input is a number or empty"""
    if value == "":
        return True
    try:
        int(value)
        return True
    except ValueError:
        return False

# ---------- App Window ----------
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

app = ctk.CTk()
app.title("Affidavit Generator - N&D Co.")

# Set application icon (logo) - works for window and taskbar
try:
    app.iconbitmap("logo.ico")
except Exception as e:
    print(f"Could not load icon: {e}")

# Optimize performance
app.update_idletasks()  # Process pending events before building UI

# Header
header_frame = ctk.CTkFrame(app, fg_color=("gray85", "gray20"))
header_frame.pack(fill="x", padx=10, pady=10)

title_label = ctk.CTkLabel(
    header_frame,
    text="🏛️ General Affidavit Generator",
    font=("Bookman Old Style", 35, "bold"),
    text_color=("#1f538d", "#63b4ff")
)
title_label.pack(pady=18)

subtitle_label = ctk.CTkLabel(
    header_frame,
    text="Generate citizenship affidavits with ease",
    font=("Bookman Old Style", 17, "italic"),
    text_color="gray"
)
subtitle_label.pack(pady=(0, 12))

# Main content frame with scrolling - Optimized settings
main_container = ctk.CTkFrame(app, fg_color="transparent")
main_container.pack(fill="both", expand=True, padx=10, pady=5)

# Use Canvas-based scrolling with better performance
frame = ctk.CTkScrollableFrame(
    main_container, 
    fg_color=("gray90", "gray15"),
    scrollbar_button_color=("gray70", "gray30"),
    scrollbar_button_hover_color=("gray60", "gray40")
)
frame.pack(fill="both", expand=True)

# Configure for smooth scrolling
frame._parent_canvas.configure(yscrollincrement=20)  # Smooth scroll increment

entries = {}

# Variable to track "same address" checkbox state
same_address_var = ctk.BooleanVar(value=False)

# Helper function to create field rows efficiently
def create_field_row(parent, label_text, field_key, padx=30, pady=8, label_width=280, height=46, font_size=16):
    """Create a field row with label and entry - optimized for large screens"""
    field_frame = ctk.CTkFrame(parent, fg_color="transparent")
    field_frame.pack(pady=pady, padx=padx, fill="x")

    label = ctk.CTkLabel(
        field_frame,
        text=label_text + " *",
        font=("Bookman Old Style", 17 if padx <= 30 else 16, "bold" if padx <= 30 else "normal"),
        anchor="w",
        width=label_width if padx <= 30 else 250
    )
    label.pack(side="left", padx=(0, 18))

    entry = ctk.CTkEntry(field_frame, height=height, font=("Bookman Old Style", font_size))
    entry.pack(side="left", fill="x", expand=True)
    # Bind focus event to auto-scroll when widget receives focus
    entry.bind('<FocusIn>', on_widget_focus)
    # Bind Tab key to custom handler that also scrolls
    entry.bind('<Tab>', focus_next_widget)
    entries[field_key] = entry
    return entry

# Section 1: Applicant Information
section1_label = ctk.CTkLabel(
    frame,
    text="📋 Applicant Information",
    font=("Bookman Old Style", 25, "bold"),
    anchor="w"
)
section1_label.pack(pady=(25, 12), padx=20, fill="x")

ctk.CTkFrame(frame, height=2, fg_color=("gray60", "gray40")).pack(fill="x", padx=20)

# Basic applicant fields (larger size for big screens)
create_field_row(frame, "Applicant Name", "Applicant Name", padx=30, pady=12, label_width=280, height=52, font_size=18)
create_field_row(frame, "Applicants Father Name", "Applicants Father Name", padx=30, pady=12, label_width=280, height=52, font_size=18)

# Indian Address subsection
indian_address_label = ctk.CTkLabel(
    frame,
    text="🏠 Indian Address",
    font=("Bookman Old Style", 20, "bold"),
    anchor="w",
    text_color=("#1f538d", "#63b4ff")
)
indian_address_label.pack(pady=(18, 6), padx=40, fill="x")

# Create Indian address fields efficiently
for field_key, field_label in [
    ("IND_Village", "Village"),
    ("IND_Post_Office", "Post Office"),
    ("IND_Police_Station", "Police Station"),
    ("IND_District", "District"),
    ("IND_Pin_Code", "Pin Code")
]:
    entry = create_field_row(frame, field_label, field_key, padx=50, pady=7, height=46, font_size=16)
    # Add trace to update introducer address if checkbox is checked
    def make_trace(ind_key, int_key):
        def on_ind_field_change(*_):
            if same_address_var.get():
                ind_entry = entries.get(ind_key)
                int_entry = entries.get(int_key)
                if ind_entry and int_entry:
                    int_entry.delete(0, 'end')
                    int_entry.insert(0, ind_entry.get().strip())
        return on_ind_field_change
    entry.bind('<KeyRelease>', make_trace(field_key, field_key.replace("IND_", "INT_")))

# Bangladesh Address subsection
bangladesh_address_label = ctk.CTkLabel(
    frame,
    text="🏠 Bangladesh Address",
    font=("Bookman Old Style", 20, "bold"),
    anchor="w",
    text_color=("#1f538d", "#63b4ff")
)
bangladesh_address_label.pack(pady=(18, 6), padx=40, fill="x")

# Create Bangladesh address fields efficiently
for field_key, field_label in [
    ("BD_Village", "Village"),
    ("BD_Post_Office", "Post Office"),
    ("BD_Police_Station", "Police Station"),
    ("BD_District", "District")
]:
    create_field_row(frame, field_label, field_key, padx=50, pady=7, height=46, font_size=16)

# Date of Entry field
date_frame = ctk.CTkFrame(frame, fg_color="transparent")
date_frame.pack(pady=12, padx=30, fill="x")

date_label = ctk.CTkLabel(
    date_frame,
    text="Date of Entry (India) *",
    font=("Bookman Old Style", 18, "bold"),
    anchor="w",
    width=280
)
date_label.pack(side="left", padx=(0, 18))

date_entry_widget = DateEntry(
    date_frame,
    font=("Bookman Old Style", 16),
    date_pattern="dd-mm-yyyy",
    width=25,
    background='darkblue',
    foreground='white',
    borderwidth=2,
    year=2020
)
date_entry_widget.pack(side="left")
# Bind events for auto-scroll
date_entry_widget.bind('<FocusIn>', on_widget_focus)
date_entry_widget.bind('<Tab>', focus_next_widget)

# Section 2: Introducer Information
section2_label = ctk.CTkLabel(
    frame,
    text="👤 Introducer/Character Reference",
    font=("Bookman Old Style", 25, "bold"),
    anchor="w"
)
section2_label.pack(pady=(35, 12), padx=20, fill="x")

ctk.CTkFrame(frame, height=2, fg_color=("gray60", "gray40")).pack(fill="x", padx=20)

# Create introducer fields efficiently (larger for big screens)
create_field_row(frame, "Introducer Name", "Introducer Name", padx=30, pady=12, label_width=280, height=52, font_size=18)
create_field_row(frame, "Introducer Occupation", "Introducer Occupation", padx=30, pady=12, label_width=280, height=52, font_size=18)
create_field_row(frame, "Introducer Father Name", "Introducer Father Name", padx=30, pady=12, label_width=280, height=52, font_size=18)

# Introducer Address subsection
introducer_address_label = ctk.CTkLabel(
    frame,
    text="🏠 Introducer Address",
    font=("Bookman Old Style", 20, "bold"),
    anchor="w",
    text_color=("#1f538d", "#63b4ff")
)
introducer_address_label.pack(pady=(18, 6), padx=40, fill="x")

# Same address checkbox
same_address_checkbox = ctk.CTkCheckBox(
    frame,
    text="Same as Applicant's Indian Address",
    variable=same_address_var,
    command=toggle_same_address,
    font=("Bookman Old Style", 16),
    checkbox_width=26,
    checkbox_height=26
)
same_address_checkbox.pack(pady=(6, 12), padx=50, anchor="w")

# Create introducer address fields efficiently
for field_key, field_label in [
    ("INT_Village", "Village"),
    ("INT_Post_Office", "Post Office"),
    ("INT_Police_Station", "Police Station"),
    ("INT_District", "District"),
    ("INT_Pin_Code", "Pin Code")
]:
    entry = create_field_row(frame, field_label, field_key, padx=50, pady=7, height=46, font_size=16)
    introducer_address_entries.append(entry)

# Age field with number input
age_frame = ctk.CTkFrame(frame, fg_color="transparent")
age_frame.pack(pady=12, padx=30, fill="x")

age_label = ctk.CTkLabel(
    age_frame,
    text="Introducer Age *",
    font=("Bookman Old Style", 18, "bold"),
    anchor="w",
    width=280
)
age_label.pack(side="left", padx=(0, 18))

# Create validation command for number-only input
vcmd = (app.register(validate_number_input), '%P')

age_entry = ctk.CTkEntry(
    age_frame,
    height=52,
    font=("Bookman Old Style", 16),
    width=150,
    validate='key',
    validatecommand=vcmd,
    placeholder_text="18-100"
)
age_entry.pack(side="left")
# Bind events for auto-scroll
age_entry.bind('<FocusIn>', on_widget_focus)
age_entry.bind('<Tab>', focus_next_widget)

age_hint = ctk.CTkLabel(
    age_frame,
    text="(Age 18-100)",
    font=("Bookman Old Style", 15, "italic"),
    text_color="gray"
)
age_hint.pack(side="left", padx=12)

# Info section - Simplified for performance
info_frame = ctk.CTkFrame(frame, fg_color=("lightblue", "gray25"), border_width=1, border_color=("gray60", "gray40"))
info_frame.pack(pady=25, padx=30, fill="x")

info_text = ctk.CTkLabel(
    info_frame,
    text="ℹ️ Current date will be automatically added to the documents",
    font=("Bookman Old Style", 16),
    text_color="gray"
)
info_text.pack(pady=12)

# Force UI update to ensure smooth rendering
frame.update_idletasks()

# Bottom frame for buttons
button_frame = ctk.CTkFrame(app, fg_color="transparent")
button_frame.pack(fill="x", padx=10, pady=18)

# Buttons (larger for big screens)
clear_btn = ctk.CTkButton(
    button_frame,
    text="🔄 Clear Form",
    command=clear_form,
    font=("Bookman Old Style", 18, "bold"),
    height=63,
    width=230,
    fg_color="gray",
    hover_color="darkgray"
)
clear_btn.pack(side="left", padx=15)

submit_btn = ctk.CTkButton(
    button_frame,
    text="📄 Generate Affidavits",
    command=submit_form,
    font=("Bookman Old Style", 18, "bold"),
    height=63,
    width=320,
    fg_color=("#1f538d", "#2980b9"),
    hover_color=("#16406b", "#1c5a8a")
)
submit_btn.pack(side="right", padx=18)

# Status label
status_label = ctk.CTkLabel(
    app,
    text="",
    font=("Bookman Old Style", 16, "italic")
)
status_label.pack(pady=10)

# Footer
footer_frame = ctk.CTkFrame(app, fg_color="transparent", height=40)
footer_frame.pack(side="bottom", fill="x", pady=10)

footer = ctk.CTkLabel(
    footer_frame,
    text="Powered by N&D Co. | www.ndcompany.in | Version 1.1",
    font=("Times New Roman", 15, "italic"),
    text_color="gray"
)
footer.pack()

# Make window fullscreen - get screen dimensions and maximize
app.update()  # Update window to ensure proper geometry calculation
screen_width = app.winfo_screenwidth()
screen_height = app.winfo_screenheight()
app.geometry(f"{screen_width}x{screen_height}+0+0")
app.state('zoomed')  # Also set zoomed state for Windows

app.mainloop()
